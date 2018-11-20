import os
import re
import xlrd
from docx import Document
from xlutils.copy import copy
'''
    参考答案.docx 与 成绩单.xls 需要同该 1612839_陈戌_2nd.py 文件位于同一目录下;
    答题卡都应位于 答题卡文件夹 中, 此文件夹也应同 1612839_陈戌_2nd.py 文件位于同一目录:
    │  1612839_陈戌_2nd.py
    │  参考答案.docx
    │  成绩单.xls
    └──答题卡
            1-罗.docx
            2-段.docx
            3-曾.docx
            4-朱.docx
            5-赵.docx
'''


def equal(rig, ans):
    rig = rig.replace(', ', '')
    res = re.search('(\w+)\((.+)\)', rig)
    if res:
        rig = re.sub('\(|\)', '', rig)
        rig += '/' + res.group(1)
    return ans.strip().lower() in [r.strip().lower() for r in rig.split('/')]


def judge(answers, inp):
    grade = [
        answers[1][i] if equal(*grade) else 0
        for i, grade in enumerate(zip(answers[0], inp[2]))
    ]
    return sum(grade[:20]), sum(grade[20:45]), sum(grade[45:50]), 0, sum(grade)


def scan(doc):
    name, id, grades = '', '', []
    for table in doc.tables[1:3]:
        for row in table.rows[1::2]:
            grades.extend([x.text for x in row.cells])
    for table in doc.tables[3:5]:
        for row in table.rows:
            grades.append(row.cells[1].text)
    for para in doc.paragraphs:
        para = re.match(r'姓名：\s+(.+)\s+学号：\s+(\d+)', para.text)
        if para:
            name, id = para.group(1), para.group(2)
    return name.strip(), id, grades


if __name__ == '__main__':
    current = os.path.abspath('./1612839_陈戌_2nd.py')[:-17]
    docs = os.listdir(current + '答题卡')
    excel = copy(xlrd.open_workbook(current + '成绩单.xls'))
    answer, grades = [], []
    for para in Document(current + '参考答案.docx').paragraphs:
        para = re.sub('参考答案|Part|Section [A-Z]|Task [1-9]|Ⅰ|Ⅱ|Ⅲ', ' ',
                      para.text).strip()
        answer.extend([x.strip() for x in re.split('\s*\d+\.\s*', para) if x])
        answers = [
            answer, [1 for i in range(20)] + [2 for i in range(29)] + [7, 15]
        ]
    for i, doc in enumerate(docs, start=1):
        path = current + '答题卡\\' + doc
        doc = Document(path)
        for j, c in zip(
                judge(answers, scan(doc)), doc.tables[0].rows[1].cells[1:]):
            c.text = str(j)
        print(doc.tables[0].rows[1].cells[-1].text)
        excel.get_sheet(0).write(i, 2, doc.tables[0].rows[1].cells[-1].text)
        doc.save(path)
    excel.save(current + '成绩单.xls')
